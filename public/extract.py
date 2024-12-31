import os
from bs4 import BeautifulSoup
from docx import Document

# Directory containing your HTML files
directory = r"C:\Users\UShahid\OneDrive\Desktop\Sanam-website-new\public"

# Output Word file
output_file = "extracted_content_english.docx"

# Tags to extract meaningful content
tags_to_extract = ["title", "p", "a", "h1", "h2", "h3", "h4", "h5", "h6", "li", "table"]

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading("Extracted Content from HTML Files", level=1)

# Walk through the directory and process files that do NOT start with 'ar_'
for root, _, files in os.walk(directory):
    for file in files:
        if not file.startswith("ar_") and file.endswith(".html"):  # Exclude files starting with 'ar_'
            file_path = os.path.join(root, file)
            with open(file_path, "r", encoding="utf-8") as infile:
                content = infile.read()
                soup = BeautifulSoup(content, "html.parser")

                # Remove the navbar (header tag with id 'header')
                header = soup.find("header", {"id": "header"})
                if header:
                    header.decompose()

                # Remove the footer (footer tag with class 'footer')
                footer = soup.find("footer", {"class": "footer"})
                if footer:
                    footer.decompose()

                # Remove styles and scripts
                for script_or_style in soup(["script", "style"]):
                    script_or_style.decompose()

                # Add a heading for each file
                doc.add_heading(file, level=2)

                # Extract content from specified tags
                for tag in tags_to_extract:
                    for element in soup.find_all(tag):
                        if tag == "table":
                            # Handle table content
                            rows = element.find_all("tr")
                            table_text = ""
                            for row in rows:
                                cells = row.find_all(["td", "th"])
                                row_text = " | ".join(cell.get_text(strip=True) for cell in cells)
                                table_text += row_text + "\n"
                            doc.add_paragraph(table_text)
                        else:
                            doc.add_paragraph(element.get_text(strip=True))

                # Add a blank line between files
                doc.add_paragraph("\n")

# Save the Word document
doc.save(output_file)

print(f"Meaningful content extracted and saved to {output_file}")

